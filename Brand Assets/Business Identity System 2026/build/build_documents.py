from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
LOGO_LIGHT = ROOT / 'build' / 'logo-raster' / 'logo-light.png'
CARBON = '2D2926'; IVORY = 'F3F0E8'; RED = 'E03C31'; GREY = '716D68'; LINE = 'D6D2C4'

def force_font(obj, name):
    """Populate every Word font slot, including Vietnamese and complex scripts."""
    obj.font.name = name
    rpr = obj._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts'); rpr.insert(0, rfonts)
    for slot in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn(f'w:{slot}'), name)
    for theme_slot in ('asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme'):
        key = qn(f'w:{theme_slot}')
        if key in rfonts.attrib:
            del rfonts.attrib[key]

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def borders(table, color=LINE, size='6'):
    tblPr=table._tbl.tblPr; e=tblPr.first_child_found_in('w:tblBorders')
    if e is None: e=OxmlElement('w:tblBorders'); tblPr.append(e)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        n=OxmlElement('w:'+edge); n.set(qn('w:val'),'single'); n.set(qn('w:sz'),size); n.set(qn('w:color'),color); e.append(n)

def margins(cell, top=100, start=120, bottom=100, end=120):
    tc=cell._tc.get_or_add_tcPr(); m=tc.first_child_found_in('w:tcMar')
    if m is None: m=OxmlElement('w:tcMar'); tc.append(m)
    for tag,val in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        x=OxmlElement('w:'+tag); x.set(qn('w:w'),str(val)); x.set(qn('w:type'),'dxa'); m.append(x)

def set_cell_text(cell, text, size=9, bold=False, color=CARBON, font='Manrope'):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(text); r.bold=bold; force_font(r,font); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color)
    margins(cell)

def setup(doc):
    sec=doc.sections[0]; sec.page_width=Mm(210); sec.page_height=Mm(297); sec.top_margin=Mm(22); sec.bottom_margin=Mm(20); sec.left_margin=Mm(23); sec.right_margin=Mm(23); sec.header_distance=Mm(9); sec.footer_distance=Mm(10)
    styles=doc.styles
    for s in styles:
        if getattr(s, 'type', None) is not None:
            try: force_font(s,'Manrope')
            except (AttributeError, ValueError): pass
    normal=styles['Normal']; force_font(normal,'Manrope'); normal.font.size=Pt(10); normal.font.color.rgb=RGBColor.from_string(CARBON); normal.paragraph_format.space_after=Pt(7); normal.paragraph_format.line_spacing=1.18
    for name,size,color,space in [('Title',34,CARBON,14),('Heading 1',20,CARBON,9),('Heading 2',13,RED,5),('Heading 3',10,CARBON,3)]:
        s=styles[name]; force_font(s,'Cormorant Garamond' if name in ('Title','Heading 1') else 'Manrope'); s.font.size=Pt(size); s.font.bold=(name not in ('Title','Heading 1')); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(space); s.paragraph_format.space_after=Pt(space/2)
    header=sec.header; p=header.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(LOGO_LIGHT), width=Mm(43))
    r=p.add_run('   HUMAN JUDGEMENT / BUSINESS IMPACT'); force_font(r,'Manrope'); r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(GREY)
    pb=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'8'); bottom.set(qn('w:color'),RED); pb.append(bottom); p._p.get_or_add_pPr().append(pb)
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    rr=fp.add_run('brandhere.co  ·  Ho Chi Minh City, Vietnam  ·  tuan.nguyen@brandhere.vn  ·  +84 907 255 734'); force_font(rr,'Manrope'); rr.font.size=Pt(7); rr.font.color.rgb=RGBColor.from_string(GREY)
    return doc

def label(doc,text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(5)
    r=p.add_run(text.upper()); force_font(r,'Manrope'); r.font.size=Pt(7); r.font.color.rgb=RGBColor.from_string(RED); r.font.bold=True

def build_letterhead():
    doc=setup(Document())
    label(doc,'Correspondence / 2026')
    p=doc.add_paragraph(style='Title'); p.add_run('[Document title]')
    meta=doc.add_table(rows=4, cols=2); meta.alignment=WD_TABLE_ALIGNMENT.LEFT; meta.autofit=False
    widths=[Mm(35),Mm(125)]
    for row in meta.rows:
        for i,c in enumerate(row.cells): c.width=widths[i]
    fields=[('DATE','[DD MONTH YYYY]'),('TO','[Recipient name · title · company]'),('FROM','Alton Nguyen · Founder & Strategic Advisor'),('SUBJECT','[Clear subject line]')]
    for row,(a,b) in zip(meta.rows,fields): set_cell_text(row.cells[0],a,7,True,RED,'Manrope'); set_cell_text(row.cells[1],b,9,False,CARBON)
    borders(meta)
    doc.add_paragraph('Dear [Name],')
    doc.add_paragraph('[Begin correspondence here. Use short paragraphs, clear decisions and specific next steps. This template is designed for client letters, formal confirmations, cover notes and business correspondence.]')
    doc.add_paragraph('[Continue body text. Keep the document concise whenever possible.]')
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(18); p.add_run('Sincerely,').italic=True
    p=doc.add_paragraph(); p.add_run('Alton Nguyen').bold=True; p.add_run('\nFounder & Strategic Advisor\nBrand Here')
    doc.add_paragraph('\n[Signature / approval area]')
    out=ROOT/'01 Stationery'/'Brand-Here-Letterhead-Concept-3.docx'; doc.save(out)

def add_scope_table(doc):
    t=doc.add_table(rows=1, cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    headers=['PHASE','WHAT WE DO','DELIVERABLE','TIMING']
    for i,h in enumerate(headers): shade(t.rows[0].cells[i],CARBON); set_cell_text(t.rows[0].cells[i],h,7,True,IVORY,'Manrope')
    data=[('[01]','[Strategic diagnosis / discovery]','[Decision brief]','[X weeks]'),('[02]','[Design / build / test]','[Working output]','[X weeks]'),('[03]','[Implementation / enablement]','[Final system + handover]','[X weeks]')]
    for vals in data:
        cells=t.add_row().cells
        for i,v in enumerate(vals): set_cell_text(cells[i],v,8,False,CARBON)
    borders(t)

def build_proposal():
    doc=setup(Document())
    label(doc,'Client proposal / statement of work')
    p=doc.add_paragraph(style='Title'); p.add_run('[One decisive outcome.]')
    p=doc.add_paragraph(); r=p.add_run('Prepared for [Client] · [Project] · [Date]'); force_font(r,'Manrope'); r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GREY)
    lead=doc.add_table(rows=1,cols=1); shade(lead.cell(0,0),CARBON); set_cell_text(lead.cell(0,0),'AI can accelerate the work. Human judgement gives it direction.',14,False,IVORY,'Cormorant Garamond'); margins(lead.cell(0,0),260,260,260,260)
    doc.add_heading('1. The decision in front of us',1); doc.add_paragraph('[Describe the business decision, not only the communication request. What must become clearer, faster or more valuable?]')
    doc.add_heading('2. What success will look like',1); doc.add_paragraph('[Define observable outcomes, decision criteria and measures of progress. Avoid vague language.]')
    doc.add_heading('3. Recommended engagement',1)
    doc.add_paragraph('Select the relevant Brand Here capability: Strategy Consulting · AI Transformation · Executive Branding · Commerce & Market Expansion.')
    add_scope_table(doc)
    doc.add_heading('4. Working model',1)
    for title,text in [('Senior judgement','Brand Here leads the problem definition, strategic direction and quality standard.'),('Expert network','Specialists join where deeper governance, finance, technology or commerce expertise materially improves the decision.'),('Partner-enabled delivery','Production and technology partners are scoped transparently around the assignment.')]:
        doc.add_heading(title,2); doc.add_paragraph(text)
    doc.add_heading('5. Investment',1)
    t=doc.add_table(rows=1,cols=3); heads=['FEE LAYER','BASIS','INVESTMENT']
    for i,h in enumerate(heads): shade(t.rows[0].cells[i],CARBON); set_cell_text(t.rows[0].cells[i],h,7,True,IVORY,'Manrope')
    for vals in [('Brand Here consulting fee','Strategy, direction and accountable lead','[VND]'),('Expert fee','Specialist advisory where approved','[VND]'),('Third-party cost','Production, technology, media or other approved suppliers','[VND]')]:
        cells=t.add_row().cells
        for i,v in enumerate(vals): set_cell_text(cells[i],v,8)
    borders(t)
    doc.add_paragraph('All third-party costs are approved before commitment. VAT and payment terms are stated in the accompanying quotation.')
    doc.add_heading('6. Assumptions and boundaries',1)
    for x in ['Two consolidated client feedback rounds per defined deliverable.','Client supplies timely access to stakeholders, data and existing materials.','Usage rights, personal data handling and use of generative AI are agreed before production.','Changes outside approved scope are quoted separately before work begins.']:
        doc.add_paragraph(x,style='List Bullet')
    doc.add_heading('7. Next step',1); doc.add_paragraph('Approve the proposed direction, confirm the commercial quotation and schedule the kickoff decision session.')
    doc.add_paragraph('\nACCEPTED FOR [CLIENT]\n\nName: ____________________    Title: ____________________\n\nSignature: _________________   Date: ____________________')
    out=ROOT/'02 Commercial Templates'/'Brand-Here-Proposal-SOW-Concept-3.docx'; doc.save(out)

build_letterhead(); build_proposal()
